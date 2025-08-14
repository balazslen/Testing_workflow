from docx import Document

def create_word_template():
    doc = Document()

    # Heading
    doc.add_heading('Executive Summary – Merrilton Robotics Communication Essentials', level=1)

    # Paragraphs
    paragraphs = [
        "At Merrilton Robotics, effective communication is central to our ability to innovate and lead in the field of autonomous systems and humanoid robotics. Clear and intentional communication ensures that our vision, values, and technical excellence are consistently reflected in every interaction—both internally and externally. In this session, we introduced two organizational communication strategies and two primary communication methods that new employees will rely on to collaborate effectively and represent Merrilton Robotics to the world.",

        "Our first communication strategy, the Direct Strategy, focuses on presenting the main message first, followed by supporting details. It is clear, concise, and efficient, making it highly effective when the audience’s reception is expected to be positive or neutral. This approach is particularly useful for urgent updates, operational announcements, and decision-making meetings where speed and clarity are paramount.",

        "The second strategy, the Indirect Strategy, begins with background and supporting details before introducing the main message. This method builds trust, softens the impact of sensitive information, and reduces potential resistance by allowing time for the audience to process context. It is especially effective when delivering negative or potentially unwelcome news in a way that maintains relationships and fosters understanding.",

        "We also reviewed two key communication methods. Email remains the most widely used formal communication tool in our organization, valued for its ability to deliver consistent, dependable information while maintaining a permanent record. It allows us to share detailed instructions, approvals, and updates efficiently across the organization. Video conferencing, on the other hand, is essential for real-time collaboration, enabling instant interaction across locations while preserving visual cues that strengthen understanding and engagement.",

        "In summary, the choice of communication strategy and method significantly affects how messages are received and acted upon at Merrilton Robotics. By mastering both direct and indirect strategies, and by leveraging email for formal documentation alongside video conferencing for interactive engagement, new employees will be equipped to contribute effectively to our collaborative culture and uphold the high standards that define our global reputation."
    ]

    for para in paragraphs:
        doc.add_paragraph(para)
        doc.add_paragraph("")  # blank line between paragraphs

    # Save the document
    filename = "Executive_Summary_Merrilton.docx"
    doc.save(filename)
    print(f"✅ Word document created and saved as '{filename}'")

if __name__ == "__main__":
    create_word_template()


##run for terminal: /home/codespace/.python/current/bin/python generate_docx_template.py
